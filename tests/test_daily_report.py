from pathlib import Path
import re
from tempfile import TemporaryDirectory
from unittest import TestCase
from unittest.mock import patch

from docx import Document

from app.services.docx_generator import DAILY_REPORT_TEMPLATE, generate_daily_report_docx


class TrendComputationTests(TestCase):
    def setUp(self):
        # Replicate the trend computation logic from the router
        self.numeric_fields = [
            "waf_attacks", "waf_blocked", "waf_ips_banned",
            "cfw_attacks", "cfw_unblocked",
            "hss_alerts", "ddos_cleanings", "ddos_blackholes", "secmaster_alerts",
            "waf_detail_attacks", "waf_detail_blocked", "waf_qps_peak_value",
            "cfw_detail_attacks", "cfw_detail_unblocked",
            "hss_detail_total", "hss_detail_fatal", "hss_detail_high",
            "hss_detail_medium", "hss_detail_low",
            "hss_unclosed_event_count",
            "ddos_detail_cleanings", "ddos_detail_blackholes",
            "secmaster_detail_total", "secmaster_detail_fatal", "secmaster_detail_high",
            "secmaster_detail_medium", "secmaster_detail_low", "secmaster_detail_info",
            "secmaster_unclosed_event_count",
        ]

    def _compute_trends(self, current, previous):
        if previous is None:
            return {}
        trends = {}
        for field in self.numeric_fields:
            cur_val = int(current.get(field, 0))
            prev_val = int(previous.get(field, 0))
            change = cur_val - prev_val
            pct = round((change / prev_val) * 100, 1) if prev_val != 0 else None
            trends[field] = {"change": change, "percent": pct}
        return trends

    def test_no_previous_returns_empty(self):
        trends = self._compute_trends({"waf_attacks": 100}, None)
        self.assertEqual(trends, {})

    def test_increase_calculated_correctly(self):
        trends = self._compute_trends(
            {"waf_attacks": 150},
            {"waf_attacks": 100},
        )
        self.assertEqual(trends["waf_attacks"]["change"], 50)
        self.assertEqual(trends["waf_attacks"]["percent"], 50.0)

    def test_decrease_calculated_correctly(self):
        trends = self._compute_trends(
            {"waf_attacks": 80},
            {"waf_attacks": 100},
        )
        self.assertEqual(trends["waf_attacks"]["change"], -20)
        self.assertEqual(trends["waf_attacks"]["percent"], -20.0)

    def test_zero_previous_returns_none_percent(self):
        trends = self._compute_trends(
            {"waf_attacks": 10},
            {"waf_attacks": 0},
        )
        self.assertEqual(trends["waf_attacks"]["change"], 10)
        self.assertIsNone(trends["waf_attacks"]["percent"])

    def test_missing_fields_treated_as_zero(self):
        trends = self._compute_trends(
            {"waf_attacks": 50},
            {},
        )
        self.assertEqual(trends["waf_attacks"]["change"], 50)
        self.assertIsNone(trends["waf_attacks"]["percent"])


class DocxGenerationTests(TestCase):
    def _normalized_xml(self, element) -> str:
        if element is None:
            return ""
        return re.sub(r">\s+<", "><", element.xml).strip()

    def _paragraph_snapshot(self, paragraph):
        paragraph_format = paragraph.paragraph_format
        first_run = paragraph.runs[0] if paragraph.runs else None
        return {
            "style": paragraph.style.name if paragraph.style else None,
            "alignment": paragraph.alignment,
            "left_indent": paragraph_format.left_indent.pt if paragraph_format.left_indent else None,
            "first_line_indent": paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else None,
            "space_before": paragraph_format.space_before.pt if paragraph_format.space_before else None,
            "space_after": paragraph_format.space_after.pt if paragraph_format.space_after else None,
            "font_name": first_run.font.name if first_run else None,
            "font_size": first_run.font.size.pt if first_run and first_run.font.size else None,
            "bold": first_run.bold if first_run else None,
        }

    def _paragraph_style_signature(self, paragraph):
        paragraph_format = paragraph.paragraph_format
        runs = []
        for run in paragraph.runs:
            if run.text or len(paragraph.runs) == 1:
                runs.append(
                    {
                        "font_name": (run.font.name or "").replace(" ", ""),
                        "font_size": run.font.size.pt if run.font.size else None,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                    }
                )
        return {
            "style": paragraph.style.name if paragraph.style else None,
            "alignment": paragraph.alignment,
            "left_indent": paragraph_format.left_indent.pt if paragraph_format.left_indent else None,
            "first_line_indent": paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else None,
            "space_before": paragraph_format.space_before.pt if paragraph_format.space_before else None,
            "space_after": paragraph_format.space_after.pt if paragraph_format.space_after else None,
            "line_spacing": paragraph_format.line_spacing,
            "runs": runs,
        }

    def _paragraph_structure_signature(self, paragraph):
        paragraph_format = paragraph.paragraph_format
        return {
            "style": paragraph.style.name if paragraph.style else None,
            "alignment": paragraph.alignment,
            "left_indent": paragraph_format.left_indent.pt if paragraph_format.left_indent else None,
            "right_indent": paragraph_format.right_indent.pt if paragraph_format.right_indent else None,
            "first_line_indent": paragraph_format.first_line_indent.pt if paragraph_format.first_line_indent else None,
            "space_before": paragraph_format.space_before.pt if paragraph_format.space_before else None,
            "space_after": paragraph_format.space_after.pt if paragraph_format.space_after else None,
            "line_spacing": paragraph_format.line_spacing,
            "line_spacing_rule": str(paragraph_format.line_spacing_rule),
        }

    def _cell_layout_signatures(self, cell):
        return [self._paragraph_structure_signature(paragraph) for paragraph in cell.paragraphs]

    def _run_format_signature(self, run):
        return {
            "rpr": self._normalized_xml(run._element.rPr),
            "drawing_count": len(run._element.xpath('.//*[local-name()="drawing"]')),
            "br_count": len(run._element.xpath('.//*[local-name()="br"]')),
            "tab_count": len(run._element.xpath('.//*[local-name()="tab"]')),
        }

    def _paragraph_format_signature(self, paragraph):
        return {
            "ppr": self._normalized_xml(paragraph._element.pPr),
            "runs": [self._run_format_signature(run) for run in paragraph.runs],
        }

    def _cell_format_signature(self, cell):
        return {
            "tcpr": self._normalized_xml(cell._tc.tcPr),
            "paragraphs": [self._paragraph_format_signature(paragraph) for paragraph in cell.paragraphs],
            "tables": [self._table_format_signature(table) for table in cell.tables],
        }

    def _row_format_signature(self, row):
        return {
            "trpr": self._normalized_xml(row._tr.trPr),
            "cells": [self._cell_format_signature(cell) for cell in row.cells],
        }

    def _table_format_signature(self, table):
        return {
            "tblpr": self._normalized_xml(table._tbl.tblPr),
            "tblgrid": self._normalized_xml(table._tbl.tblGrid),
            "rows": [self._row_format_signature(row) for row in table.rows],
        }

    def _section_format_signature(self, section):
        section_properties = section._sectPr
        return {
            "sectpr": self._normalized_xml(section_properties),
            "header_paragraphs": [self._paragraph_format_signature(paragraph) for paragraph in section.header.paragraphs],
            "footer_paragraphs": [self._paragraph_format_signature(paragraph) for paragraph in section.footer.paragraphs],
        }

    def _document_format_signature(self, document):
        return {
            "sections": [self._section_format_signature(section) for section in document.sections],
            "body_paragraphs": [self._paragraph_format_signature(paragraph) for paragraph in document.paragraphs],
            "tables": [self._table_format_signature(table) for table in document.tables],
        }

    def test_generates_docx_file(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "暂无",
            "key_work_content": "暂无",
        }
        operators = [
            {"name": "张三", "phone": "13800138000", "role": "安全运营人员", "responsibility": "负责安全监控。"},
        ]

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, operators)

            self.assertTrue(file_path.exists())
            self.assertEqual(file_path.suffix, ".docx")
            self.assertIn("2026-05-21", file_path.name)

    def test_generates_docx_with_screenshot_paths(self):
        png_bytes = bytes.fromhex(
            "89504E470D0A1A0A0000000D49484452000000010000000108060000001F15C489"
            "0000000D49444154789C6360606060000000050001A5F645400000000049454E44AE426082"
        )
        with TemporaryDirectory() as temp_dir:
            screenshot_dir = Path(temp_dir) / "shots"
            screenshot_dir.mkdir(parents=True, exist_ok=True)
            screenshot_path = screenshot_dir / "waf.png"
            screenshot_path.write_bytes(png_bytes)

            report = {
                "report_date": "2026-05-21",
                "waf_screenshot_path": str(screenshot_path),
                "business_stability": "稳定",
            }

            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, [])

            self.assertTrue(file_path.exists())

    def test_populates_template_fixed_sections(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }
        operators = [
            {"name": "张三", "phone": "13800138000", "role": "安全运营人员", "responsibility": "负责安全监控。"},
        ]

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, operators)

            doc = Document(file_path)
            table = doc.tables[0]

            self.assertEqual(table.rows[2].cells[0].paragraphs[0].text, "今日业务运行稳定。")
            self.assertIn("WAF遭受攻击1000次", table.rows[2].cells[0].paragraphs[1].text)
            self.assertIn("趋势平稳。总体态势良好。", table.rows[2].cells[0].paragraphs[2].text)
            self.assertIn("监控时间：2026年5月20日 18:00~2026年5月21日 18:00", table.rows[3].cells[0].paragraphs[0].text)
            self.assertIn("未闭环事件12次", table.rows[4].cells[0].paragraphs[13].text)
            self.assertEqual(table.rows[6].cells[0].paragraphs[0].text, "攻击路径分析暂无异常。")
            self.assertEqual(table.rows[8].cells[0].paragraphs[0].text, "今日已完成重点巡检。")

            operator_table = table.rows[10].cells[0].tables[0]
            self.assertEqual(len(operator_table.rows), 5)
            self.assertEqual(operator_table.rows[1].cells[1].text, "张三")
            self.assertEqual(operator_table.rows[1].cells[4].text, "负责安全监控。")
            self.assertEqual(operator_table.rows[2].cells[1].text, "")

    def test_preserves_template_paragraph_formatting(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, [])

            template = Document(DAILY_REPORT_TEMPLATE)
            generated = Document(file_path)

            paragraph_positions = [(2, 0), (2, 1), (2, 2), (3, 0), (4, 1), (6, 0)]
            for row_index, paragraph_index in paragraph_positions:
                template_paragraph = template.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                generated_paragraph = generated.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                self.assertEqual(
                    self._paragraph_snapshot(generated_paragraph),
                    self._paragraph_snapshot(template_paragraph),
                )

            key_work_paragraph = generated.tables[0].rows[8].cells[0].paragraphs[0]
            self.assertEqual((key_work_paragraph.runs[0].font.name or "").replace(" ", ""), "微软雅黑")

    def test_preserves_template_styles_for_all_replaced_sections(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }
        operators = [
            {"name": "张三", "phone": "13800138000", "role": "安全运营人员", "responsibility": "负责安全监控。"},
            {"name": "李四", "phone": "13900139000", "role": "项目PM", "responsibility": "负责安全监控。"},
        ]

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, operators)

            template = Document(DAILY_REPORT_TEMPLATE)
            generated = Document(file_path)

            paragraph_positions = [
                (3, 0),
                (4, 0),
                (4, 1),
                (4, 4),
                (4, 7),
                (4, 10),
                (4, 13),
                (4, 17),
                (4, 20),
                (4, 23),
                (4, 24),
                (6, 0),
            ]
            for row_index, paragraph_index in paragraph_positions:
                template_paragraph = template.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                generated_paragraph = generated.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                self.assertEqual(
                    self._paragraph_style_signature(generated_paragraph),
                    self._paragraph_style_signature(template_paragraph),
                )

            template_key_work = template.tables[0].rows[8].cells[0].paragraphs[0]
            generated_key_work = generated.tables[0].rows[8].cells[0].paragraphs[0]
            self.assertEqual(
                self._paragraph_structure_signature(generated_key_work),
                self._paragraph_structure_signature(template_key_work),
            )
            for run in generated_key_work.runs:
                if run.text:
                    self.assertEqual((run.font.name or "").replace(" ", ""), "微软雅黑")

            template_operator_table = template.tables[0].rows[10].cells[0].tables[0]
            generated_operator_table = generated.tables[0].rows[10].cells[0].tables[0]
            operator_cells = [(1, 0), (1, 1), (1, 2), (1, 3), (1, 4)]
            for row_index, cell_index in operator_cells:
                template_paragraph = template_operator_table.rows[row_index].cells[cell_index].paragraphs[0]
                generated_paragraph = generated_operator_table.rows[row_index].cells[cell_index].paragraphs[0]
                self.assertEqual(
                    self._paragraph_style_signature(generated_paragraph),
                    self._paragraph_style_signature(template_paragraph),
                )

    def test_body_paragraphs_use_microsoft_yahei(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, [])

            doc = Document(file_path)
            paragraph_positions = [
                (2, 0),
                (2, 1),
                (2, 2),
                (4, 1),
                (4, 4),
                (4, 7),
                (4, 10),
                (4, 13),
                (4, 17),
                (4, 20),
                (4, 24),
                (6, 0),
                (8, 0),
            ]
            for row_index, paragraph_index in paragraph_positions:
                paragraph = doc.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                for run in paragraph.runs:
                    if run.text.strip():
                        self.assertEqual((run.font.name or "").replace(" ", ""), "微软雅黑")

    def test_figure_captions_preserve_template_formatting(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, [])

            template = Document(DAILY_REPORT_TEMPLATE)
            generated = Document(file_path)
            caption_positions = [(4, 3), (4, 6), (4, 9), (4, 12), (4, 16), (4, 19), (4, 22)]

            for row_index, paragraph_index in caption_positions:
                template_paragraph = template.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                generated_paragraph = generated.tables[0].rows[row_index].cells[0].paragraphs[paragraph_index]
                self.assertEqual(
                    self._paragraph_style_signature(generated_paragraph),
                    self._paragraph_style_signature(template_paragraph),
                )
                for run in generated_paragraph.runs:
                    if run.text.strip():
                        self.assertEqual((run.font.name or "").replace(" ", ""), "宋体")

    def test_preserves_template_layout_for_full_report_regions(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }
        operators = [
            {"name": "张三", "phone": "13800138000", "role": "安全运营人员", "responsibility": "负责安全监控。"},
            {"name": "李四", "phone": "13900139000", "role": "项目PM", "responsibility": "负责安全监控。"},
        ]

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, operators)

            template = Document(DAILY_REPORT_TEMPLATE)
            generated = Document(file_path)

            for row_index in (1, 2, 3, 4, 5, 6, 7, 8, 9, 11):
                template_cell = template.tables[0].rows[row_index].cells[0]
                generated_cell = generated.tables[0].rows[row_index].cells[0]
                self.assertEqual(
                    self._cell_layout_signatures(generated_cell),
                    self._cell_layout_signatures(template_cell),
                )

            template_operator_table = template.tables[0].rows[10].cells[0].tables[0]
            generated_operator_table = generated.tables[0].rows[10].cells[0].tables[0]
            for row_index in range(len(template_operator_table.rows)):
                for cell_index in range(len(template_operator_table.columns)):
                    template_cell = template_operator_table.rows[row_index].cells[cell_index]
                    generated_cell = generated_operator_table.rows[row_index].cells[cell_index]
                    self.assertEqual(
                        self._cell_layout_signatures(generated_cell),
                        self._cell_layout_signatures(template_cell),
                    )

    def test_preserves_full_document_format_signature(self):
        report = {
            "report_date": "2026-05-21",
            "business_stability": "今日业务运行稳定。",
            "waf_attacks": 1000,
            "waf_blocked": 900,
            "waf_ips_banned": 5,
            "cfw_attacks": 500,
            "cfw_unblocked": 2,
            "hss_alerts": 100,
            "ddos_cleanings": 0,
            "ddos_blackholes": 0,
            "secmaster_alerts": 50,
            "trend_comparison": "趋势平稳。",
            "overall_assessment": "总体态势良好。",
            "monitor_start": "2026-05-20 18:00",
            "monitor_end": "2026-05-21 18:00",
            "waf_detail_attacks": 1000,
            "waf_detail_blocked": 900,
            "waf_qps_specs": "85,000",
            "waf_qps_peak_range": "17:00-18:00",
            "waf_qps_peak_value": 178848,
            "waf_exceeded_spec": 1,
            "cfw_detail_attacks": 500,
            "cfw_detail_unblocked": 2,
            "cfw_bandwidth_spec": "12050Mbps",
            "cfw_peak_inbound_range": "16:44-18:00",
            "cfw_inbound_peak": "15.90Gbps",
            "cfw_inbound_95th": "14.71Gbps",
            "cfw_exceeded_spec": 1,
            "hss_detail_total": 100,
            "hss_detail_fatal": 0,
            "hss_detail_high": 20,
            "hss_detail_medium": 30,
            "hss_detail_low": 50,
            "hss_unclosed_event_count": 12,
            "hss_closed_loop_status": "已全部闭环。",
            "ddos_detail_cleanings": 0,
            "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 50,
            "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0,
            "secmaster_detail_medium": 50,
            "secmaster_detail_low": 0,
            "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 8,
            "emergency_response": "无。",
            "attack_path_assessment": "攻击路径分析暂无异常。",
            "key_work_content": "今日已完成重点巡检。",
        }
        operators = [
            {"name": "张三", "phone": "13800138000", "role": "安全运营人员", "responsibility": "负责安全监控。"},
            {"name": "李四", "phone": "13900139000", "role": "项目PM", "responsibility": "负责安全监控。"},
        ]

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, operators)

            template = Document(DAILY_REPORT_TEMPLATE)
            generated = Document(file_path)

            self.assertEqual(
                self._document_format_signature(generated),
                self._document_format_signature(template),
            )

    def test_handles_empty_operators(self):
        report = {
            "report_date": "2026-05-21",
            "waf_attacks": 0, "waf_blocked": 0, "waf_ips_banned": 0,
            "cfw_attacks": 0, "cfw_unblocked": 0,
            "hss_alerts": 0, "ddos_cleanings": 0, "ddos_blackholes": 0,
            "secmaster_alerts": 0,
            "waf_detail_attacks": 0, "waf_detail_blocked": 0,
            "waf_qps_specs": "", "waf_qps_peak_range": "", "waf_qps_peak_value": 0,
            "waf_exceeded_spec": 0,
            "cfw_detail_attacks": 0, "cfw_detail_unblocked": 0,
            "cfw_bandwidth_spec": "", "cfw_peak_inbound_range": "",
            "cfw_inbound_peak": "", "cfw_inbound_95th": "", "cfw_exceeded_spec": 0,
            "hss_detail_total": 0, "hss_detail_fatal": 0,
            "hss_detail_high": 0, "hss_detail_medium": 0, "hss_detail_low": 0,
            "hss_unclosed_event_count": 0,
            "hss_closed_loop_status": "",
            "ddos_detail_cleanings": 0, "ddos_detail_blackholes": 0,
            "secmaster_detail_total": 0, "secmaster_detail_fatal": 0,
            "secmaster_detail_high": 0, "secmaster_detail_medium": 0,
            "secmaster_detail_low": 0, "secmaster_detail_info": 0,
            "secmaster_unclosed_event_count": 0,
            "business_stability": "", "trend_comparison": "", "overall_assessment": "",
            "monitor_start": "", "monitor_end": "",
            "emergency_response": "", "attack_path_assessment": "", "key_work_content": "",
        }

        with TemporaryDirectory() as temp_dir:
            with patch("app.services.docx_generator.EXPORT_DIR", Path(temp_dir)):
                file_path = generate_daily_report_docx(report, [])

            self.assertTrue(file_path.exists())
