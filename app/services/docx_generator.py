from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.config import EXPORT_DIR


BASE_DIR = Path(__file__).resolve().parent.parent.parent
DAILY_REPORT_TEMPLATE = BASE_DIR / "app" / "static" / "report-templates" / "daily-report-template.docx"


def generate_daily_report_docx(report: dict, operators: list[dict]) -> Path:
    if not DAILY_REPORT_TEMPLATE.exists():
        raise FileNotFoundError(f"日报模板不存在: {DAILY_REPORT_TEMPLATE}")
    export_dir = EXPORT_DIR / "daily"
    export_dir.mkdir(parents=True, exist_ok=True)
    file_path = export_dir / f"{report.get('report_date', '')}.docx"
    _render_docx_template(DAILY_REPORT_TEMPLATE, file_path, _build_template_context(report, operators))
    return file_path


def _render_docx_template(template_path: Path, output_path: Path, context: dict[str, object]) -> None:
    replacements = {
        f"{{{{ {key} }}}}": escape(str(value or ""))
        for key, value in context.items()
    }

    with ZipFile(template_path, "r") as source_zip, ZipFile(output_path, "w", compression=ZIP_DEFLATED) as target_zip:
        for info in source_zip.infolist():
            data = source_zip.read(info.filename)
            if info.filename == "word/document.xml":
                xml_text = data.decode("utf-8")
                for placeholder, value in replacements.items():
                    xml_text = xml_text.replace(placeholder, value)
                data = xml_text.encode("utf-8")
            target_zip.writestr(info, data)


def _build_template_context(report: dict, operators: list[dict]) -> dict[str, object]:
    normalized_operators = list(operators[:4])
    while len(normalized_operators) < 4:
        normalized_operators.append({})

    shared_responsibility = ""
    for operator in operators:
        shared_responsibility = str(operator.get("responsibility", "")).strip()
        if shared_responsibility:
            break

    return {
        "business_stability": str(report.get("business_stability", "")).strip() or "暂无",
        "summary_sentence": _compose_summary_sentence(report),
        "trend_assessment": _compose_trend_assessment(report),
        "monitor_heading": _compose_monitor_heading(report),
        "monitor_subheading": "攻击告警监控：",
        "waf_detail": _compose_waf_detail(report),
        "waf_qps_detail": _compose_waf_qps_detail(report),
        "cfw_detail": _compose_cfw_detail(report),
        "cfw_bandwidth_detail": _compose_cfw_bandwidth_detail(report),
        "hss_detail": _compose_hss_detail(report),
        "ddos_detail": _compose_ddos_detail(report),
        "secmaster_detail": _compose_secmaster_detail(report),
        "emergency_heading": "事件应急响应：",
        "emergency_response": str(report.get("emergency_response", "")).strip() or "暂无",
        "attack_path_assessment": str(report.get("attack_path_assessment", "")).strip() or "暂无",
        "key_work_content": str(report.get("key_work_content", "")).strip() or "暂无",
        "operator_group": "运营人员" if any(op for op in operators) else "",
        "operator_shared_responsibility": shared_responsibility,
        "operator_1_name": str(normalized_operators[0].get("name", "")),
        "operator_1_phone": str(normalized_operators[0].get("phone", "")),
        "operator_1_role": str(normalized_operators[0].get("role", "")),
        "operator_2_name": str(normalized_operators[1].get("name", "")),
        "operator_2_phone": str(normalized_operators[1].get("phone", "")),
        "operator_2_role": str(normalized_operators[1].get("role", "")),
        "operator_3_name": str(normalized_operators[2].get("name", "")),
        "operator_3_phone": str(normalized_operators[2].get("phone", "")),
        "operator_3_role": str(normalized_operators[2].get("role", "")),
        "operator_4_name": str(normalized_operators[3].get("name", "")),
        "operator_4_phone": str(normalized_operators[3].get("phone", "")),
        "operator_4_role": str(normalized_operators[3].get("role", "")),
    }


def _set_narrow_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _format_date_display(date_str: str) -> str:
    try:
        dt = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{dt.year}年{dt.month}月{dt.day}日"
    except ValueError:
        return date_str


def _format_monitor_display(value: object) -> str:
    raw = str(value or "").strip()
    if not raw:
        return raw
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S"):
        try:
            dt = datetime.strptime(raw, fmt)
            return f"{dt.year}年{dt.month}月{dt.day}日 {dt:%H:%M}"
        except ValueError:
            continue
    return raw


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "SimHei"


def _replace_paragraph_text(paragraph, text: str, preserve_prefix_runs: int = 0, target_font_name: str | None = None) -> None:
    runs = list(paragraph.runs)
    if runs:
        prefix_runs = runs[:preserve_prefix_runs]
        target_runs = runs[preserve_prefix_runs:]
        if not target_runs:
            target_runs = [runs[-1]]

        remaining = text
        for index, run in enumerate(target_runs):
            if index == len(target_runs) - 1:
                run.text = remaining
            else:
                original_len = len(run.text)
                run.text = remaining[:original_len]
                remaining = remaining[original_len:]
            if target_font_name:
                run.font.name = target_font_name

        for run in prefix_runs:
            run.text = run.text
        return
    if text:
        run = paragraph.add_run(text)
        if target_font_name:
            run.font.name = target_font_name


def _replace_paragraph_picture(paragraph, path_value: object) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)
    screenshot_path = _resolve_report_asset(path_value)
    if screenshot_path is None:
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(screenshot_path), width=Cm(15.5))


def _compose_summary_sentence(report: dict) -> str:
    parts = [
        f"WAF遭受攻击{report.get('waf_attacks', 0)}次，拦截攻击告警{report.get('waf_blocked', 0)}次，已对{report.get('waf_ips_banned', 0)}个IP进行汇报封禁处理；",
        f"CFW遭受攻击{report.get('cfw_attacks', 0)}次，{report.get('cfw_unblocked', 0)}次攻击未被拦截；",
        f"HSS发生告警{report.get('hss_alerts', 0)}次，{_summary_unclosed_event_text(report.get('hss_unclosed_event_count', 0))}；",
        f"DDos防护清洗{report.get('ddos_cleanings', 0)}次，{'无' if int(report.get('ddos_blackholes', 0) or 0) == 0 else f"{report.get('ddos_blackholes', 0)}次"}黑洞；",
        f"SecMaster检测告警{report.get('secmaster_alerts', 0)}次，{_summary_unclosed_event_text(report.get('secmaster_unclosed_event_count', 0))}。",
    ]
    return "".join(parts)


def _compose_trend_assessment(report: dict) -> str:
    trend = str(report.get("trend_comparison", "")).strip()
    assessment = str(report.get("overall_assessment", "")).strip()
    if trend and assessment:
        return f"{trend}{assessment}"
    return trend or assessment or "暂无"


def _compose_monitor_heading(report: dict) -> str:
    start = _format_monitor_display(report.get("monitor_start", ""))
    end = _format_monitor_display(report.get("monitor_end", ""))
    if start and end:
        return f"二、安全监控（监控时间：{start}~{end}）"
    return "二、安全监控"


def _compose_waf_detail(report: dict) -> str:
    return (
        f"WAF应用防火墙遭受攻击{report.get('waf_detail_attacks', 0)}次，"
        f"拦截攻击告警{report.get('waf_detail_blocked', 0)}次，已对{report.get('waf_ips_banned', 0)}个IP进行汇报封禁处理；"
    )


def _compose_waf_qps_detail(report: dict) -> str:
    exceeded = ""
    if int(report.get("waf_exceeded_spec", 0)):
        exceeded = "已超出规格。可能出现限流、随机丢包、自动Bypass等现象，影响业务。"
    return (
        f"当前WAF产品QPS的规格为{report.get('waf_qps_specs', '')}，"
        f"监测到QPS最大值时间段为{report.get('waf_qps_peak_range', '')}，峰值为{report.get('waf_qps_peak_value', 0)}，"
        f"{exceeded}"
    ).strip("，")


def _compose_cfw_detail(report: dict) -> str:
    return f"CFW遭受攻击{report.get('cfw_detail_attacks', 0)}次攻击，{report.get('cfw_detail_unblocked', 0)}次攻击未被拦截；"


def _compose_cfw_bandwidth_detail(report: dict) -> str:
    exceeded = ""
    if int(report.get("cfw_exceeded_spec", 0)):
        exceeded = "入方向流量峰值、入方向95带宽已超出规格。可能出现限流、随机丢包、自动Bypass等现象，影响业务。"
    return (
        f"当前CFW产品互联网边界防护带宽为{report.get('cfw_bandwidth_spec', '')}。"
        f"监测到带宽峰值时间段为{report.get('cfw_peak_inbound_range', '')}，入方向流量峰值{report.get('cfw_inbound_peak', '')}，"
        f"入方向95带宽值{report.get('cfw_inbound_95th', '')}，{exceeded}"
    ).strip("，")


def _compose_hss_detail(report: dict) -> str:
    return (
        f"HSS发生告警{report.get('hss_detail_total', 0)}次，其中致命告警{report.get('hss_detail_fatal', 0)}次，"
        f"高危告警{report.get('hss_detail_high', 0)}次，中危告警{report.get('hss_detail_medium', 0)}次，"
        f"低危告警{report.get('hss_detail_low', 0)}次，未闭环事件{report.get('hss_unclosed_event_count', 0)}次，"
        f"{report.get('hss_closed_loop_status', '')}。"
    )


def _compose_ddos_detail(report: dict) -> str:
    blackholes = report.get("ddos_detail_blackholes", 0)
    blackhole_text = "无" if int(blackholes or 0) == 0 else f"{blackholes}次"
    return f"当前DDOS原生基础防护和DDOS原生高级防护存在IP清洗{report.get('ddos_detail_cleanings', 0)}次，{blackhole_text}黑洞。"


def _compose_secmaster_detail(report: dict) -> str:
    return (
        f"SecMaster发生告警{report.get('secmaster_detail_total', 0)}次，其中致命告警{report.get('secmaster_detail_fatal', 0)}次，"
        f"高危告警{report.get('secmaster_detail_high', 0)}次，中危告警{report.get('secmaster_detail_medium', 0)}次，"
        f"低危告警{report.get('secmaster_detail_low', 0)}次，提示告警{report.get('secmaster_detail_info', 0)}次，"
        f"未闭环事件{report.get('secmaster_unclosed_event_count', 0)}次。"
    )


def _fill_template_overview(cell, report: dict) -> None:
    paragraphs = cell.paragraphs
    _replace_paragraph_text(paragraphs[0], str(report.get("business_stability", "")).strip() or "暂无", target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[1], _compose_summary_sentence(report), target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[2], _compose_trend_assessment(report), target_font_name="微软雅黑")


def _fill_template_monitoring(cell, report: dict) -> None:
    paragraphs = cell.paragraphs
    _replace_paragraph_text(paragraphs[0], "攻击告警监控：", preserve_prefix_runs=2, target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[1], _compose_waf_detail(report), target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[4], _compose_waf_qps_detail(report), preserve_prefix_runs=1, target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[7], _compose_cfw_detail(report), target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[10], _compose_cfw_bandwidth_detail(report), target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[13], _compose_hss_detail(report), target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[17], _compose_ddos_detail(report), preserve_prefix_runs=1, target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[20], _compose_secmaster_detail(report), preserve_prefix_runs=1, target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[23], "事件应急响应：", preserve_prefix_runs=2, target_font_name="微软雅黑")
    _replace_paragraph_text(paragraphs[24], str(report.get("emergency_response", "")).strip() or "暂无", target_font_name="微软雅黑")


def _fill_simple_template_section(cell, text: str) -> None:
    _replace_paragraph_text(cell.paragraphs[0], text.strip() or "暂无", target_font_name="微软雅黑")


def _fill_template_operators(cell, operators: list[dict]) -> None:
    if not cell.tables:
        return
    table = cell.tables[0]
    body_rows = list(table.rows)[1:]
    rows = operators[: len(body_rows)] if operators else []
    shared_responsibility = ""
    for operator in rows:
        shared_responsibility = str(operator.get("responsibility", "")).strip()
        if shared_responsibility:
            break
    if body_rows:
        _replace_paragraph_text(body_rows[0].cells[0].paragraphs[0], "运营人员" if rows else "", target_font_name="微软雅黑")
        _replace_paragraph_text(body_rows[0].cells[4].paragraphs[0], shared_responsibility, target_font_name="微软雅黑")
    for row_index, row in enumerate(body_rows):
        operator = rows[row_index] if row_index < len(rows) else {}
        _replace_paragraph_text(row.cells[1].paragraphs[0], str(operator.get("name", "")), target_font_name="微软雅黑")
        _replace_paragraph_text(row.cells[2].paragraphs[0], str(operator.get("phone", "")), target_font_name="微软雅黑")
        _replace_paragraph_text(row.cells[3].paragraphs[0], str(operator.get("role", "")), target_font_name="微软雅黑")


def _add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = "SimHei"
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def _add_sub_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.name = "SimHei"
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def _add_para(doc: Document, text: str) -> None:
    text = str(text or "").strip()
    if not text:
        text = "暂无"
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "SimSun"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "SimHei"

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "SimSun"

    doc.add_paragraph()


def _resolve_report_asset(path_value: object) -> Path | None:
    raw = str(path_value or "").strip()
    if not raw:
        return None
    candidate = Path(raw)
    if not candidate.is_absolute():
        candidate = BASE_DIR / raw
    return candidate if candidate.exists() else None


def _add_optional_screenshot(doc: Document, path_value: object) -> None:
    screenshot_path = _resolve_report_asset(path_value)
    if screenshot_path is None:
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(screenshot_path), width=Cm(15.5))
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)


def _add_section_1(doc: Document, report: dict) -> None:
    _add_section_heading(doc, "一、总体安全态势")

    business = str(report.get("business_stability", ""))
    _add_para(doc, business)

    summary_headers = ["安全产品", "遭受攻击", "拦截/告警", "未阻断/封禁IP"]
    summary_rows = [
        ["WAF应用防火墙", str(report.get("waf_attacks", 0)), str(report.get("waf_blocked", 0)),
         f"已对{report.get('waf_ips_banned', 0)}个IP进行汇报封禁处理"],
        ["CFW云防火墙", str(report.get("cfw_attacks", 0)), "-",
         f"{report.get('cfw_unblocked', 0)}次攻击未被拦截"],
        ["HSS主机安全", "-", f"发生告警{report.get('hss_alerts', 0)}次", _summary_unclosed_event_text(report.get("hss_unclosed_event_count", 0))],
        ["DDoS防护", "-", f"清洗{report.get('ddos_cleanings', 0)}次",
         f"{'无' if int(report.get('ddos_blackholes', 0)) == 0 else report.get('ddos_blackholes', 0)}黑洞"],
        ["SecMaster态势感知", "-", f"检测告警{report.get('secmaster_alerts', 0)}次", _summary_unclosed_event_text(report.get("secmaster_unclosed_event_count", 0))],
    ]
    _add_table(doc, summary_headers, summary_rows)

    trend = str(report.get("trend_comparison", ""))
    if trend:
        _add_para(doc, trend)

    assessment = str(report.get("overall_assessment", ""))
    if assessment:
        _add_para(doc, assessment)


def _summary_unclosed_event_text(count: object) -> str:
    return f"有{int(count)}起事件未闭环" if int(count) else "无未闭环事件"


def _add_section_2(doc: Document, report: dict) -> None:
    monitor_start = str(report.get("monitor_start", ""))
    monitor_end = str(report.get("monitor_end", ""))
    heading = "二、安全监控"
    if monitor_start and monitor_end:
        heading += f"（监控时间：{monitor_start}~{monitor_end}）"
    _add_section_heading(doc, heading)

    _add_sub_heading(doc, "1. 攻击告警监控：")

    _add_waf_detail(doc, report)
    _add_cfw_detail(doc, report)
    _add_hss_detail(doc, report)
    _add_ddos_detail(doc, report)
    _add_secmaster_detail(doc, report)

    _add_sub_heading(doc, "2. 事件应急响应：")
    _add_para(doc, str(report.get("emergency_response", "")))


def _add_waf_detail(doc: Document, report: dict) -> None:
    blocked = report.get("waf_detail_blocked", 0)
    _add_para(doc,
        f"WAF应用防火墙遭受攻击{report.get('waf_detail_attacks', 0)}次，"
        f"拦截攻击告警{blocked}次，已对{report.get('waf_ips_banned', 0)}个IP进行汇报封禁处理；"
    )
    _add_optional_screenshot(doc, report.get("waf_screenshot_path", ""))
    qps_specs = str(report.get("waf_qps_specs", ""))
    peak_range = str(report.get("waf_qps_peak_range", ""))
    peak_val = report.get("waf_qps_peak_value", 0)
    if qps_specs or peak_range or peak_val:
        exceeded = ""
        if int(report.get("waf_exceeded_spec", 0)):
            exceeded = "已超出规格。可能出现限流、随机丢包、自动Bypass等现象，影响业务。"
        _add_para(doc,
            f"当前WAF产品QPS的规格为{qps_specs}，"
            f"监测到QPS最大值时间段为{peak_range}，峰值为{peak_val}。" + exceeded
        )
        _add_optional_screenshot(doc, report.get("waf_qps_screenshot_path", ""))


def _add_cfw_detail(doc: Document, report: dict) -> None:
    _add_para(doc,
        f"CFW遭受攻击{report.get('cfw_detail_attacks', 0)}次攻击，"
        f"{report.get('cfw_detail_unblocked', 0)}次攻击未被拦截；"
    )
    _add_optional_screenshot(doc, report.get("cfw_screenshot_path", ""))
    bw_spec = str(report.get("cfw_bandwidth_spec", ""))
    peak_range = str(report.get("cfw_peak_inbound_range", ""))
    inbound_peak = str(report.get("cfw_inbound_peak", ""))
    inbound_95 = str(report.get("cfw_inbound_95th", ""))
    if bw_spec or peak_range:
        exceeded = ""
        if int(report.get("cfw_exceeded_spec", 0)):
            exceeded = "入方向流量峰值、入方向95带宽已超出规格。可能出现限流、随机丢包、自动Bypass等现象，影响业务。"
        _add_para(doc,
            f"当前CFW产品互联网边界防护带宽为{bw_spec}。"
            f"监测到带宽峰值时间段为{peak_range}，"
            f"入方向流量峰值{inbound_peak}，入方向95带宽值{inbound_95}。" + exceeded
        )


def _add_hss_detail(doc: Document, report: dict) -> None:
    total = report.get("hss_detail_total", 0)
    fatal = report.get("hss_detail_fatal", 0)
    high = report.get("hss_detail_high", 0)
    medium = report.get("hss_detail_medium", 0)
    low = report.get("hss_detail_low", 0)
    unclosed_count = report.get("hss_unclosed_event_count", 0)
    closed = str(report.get("hss_closed_loop_status", ""))
    unclosed_count_text = f"，未闭环事件{unclosed_count}次"
    _add_para(doc,
        f"HSS发生告警{total}次，其中致命告警{fatal}次，高危告警{high}次，"
        f"中危告警{medium}次，低危告警{low}次{unclosed_count_text}，{closed}。"
    )
    _add_optional_screenshot(doc, report.get("hss_screenshot_path", ""))


def _add_ddos_detail(doc: Document, report: dict) -> None:
    cleanings = report.get("ddos_detail_cleanings", 0)
    blackholes = report.get("ddos_detail_blackholes", 0)
    bh_text = "无" if int(blackholes) == 0 else f"{blackholes}次"
    _add_para(doc,
        f"当前DDOS原生基础防护和DDOS原生高级防护存在IP清洗{cleanings}次，{bh_text}黑洞。"
    )
    _add_optional_screenshot(doc, report.get("ddos_screenshot_path", ""))


def _add_secmaster_detail(doc: Document, report: dict) -> None:
    total = report.get("secmaster_detail_total", 0)
    fatal = report.get("secmaster_detail_fatal", 0)
    high = report.get("secmaster_detail_high", 0)
    medium = report.get("secmaster_detail_medium", 0)
    low = report.get("secmaster_detail_low", 0)
    info = report.get("secmaster_detail_info", 0)
    unclosed_count = report.get("secmaster_unclosed_event_count", 0)
    unclosed_count_text = f"，未闭环事件{unclosed_count}次"
    _add_para(doc,
        f"SecMaster发生告警{total}次，其中致命告警{fatal}次，高危告警{high}次，"
        f"中危告警{medium}次，低危告警{low}次，提示告警{info}次{unclosed_count_text}。"
    )
    _add_optional_screenshot(doc, report.get("secmaster_screenshot_path", ""))


def _add_section_3(doc: Document, report: dict) -> None:
    _add_section_heading(doc, "三、攻击路径评估")
    _add_para(doc, str(report.get("attack_path_assessment", "")))
    _add_optional_screenshot(doc, report.get("attack_path_screenshot_path", ""))


def _add_section_4(doc: Document, report: dict) -> None:
    _add_section_heading(doc, "四、重点工作内容")
    _add_para(doc, str(report.get("key_work_content", "")))
    _add_optional_screenshot(doc, report.get("key_work_screenshot_path", ""))


def _add_section_5(doc: Document, operators: list[dict]) -> None:
    _add_section_heading(doc, "五、运营人员")

    headers = ["所属组别", "姓名", "联系方式", "角色", "工作职责"]
    rows = []
    for op in operators:
        rows.append([
            "运营人员",
            str(op.get("name", "")),
            str(op.get("phone", "")),
            str(op.get("role", "")),
            str(op.get("responsibility", "")),
        ])

    if not rows:
        rows.append(["", "", "", "", ""])

    rows.append(["华为云计算-安全专业服务交付团队", "", "", "", ""])
    _add_table(doc, headers, rows)
