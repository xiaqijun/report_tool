from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent.parent
SOURCE_DOCX = BASE_DIR / "比亚迪规划院安全运营日报-2026年5月20日.docx"
TARGET_DOCX = BASE_DIR / "app" / "static" / "report-templates" / "daily-report-template.docx"


def set_paragraph_runs(paragraph, texts: list[str], target_font_name: str | None = None) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run("".join(texts))
        runs = list(paragraph.runs)

    for index, run in enumerate(runs):
        run.text = texts[index] if index < len(texts) else ""
        if target_font_name and run.text:
            run.font.name = target_font_name


def main() -> None:
    TARGET_DOCX.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(SOURCE_DOCX)
    table = doc.tables[0]

    set_paragraph_runs(table.rows[2].cells[0].paragraphs[0], ["{{ business_stability }}"], target_font_name="微软雅黑")
    set_paragraph_runs(table.rows[2].cells[0].paragraphs[1], ["{{ summary_sentence }}"], target_font_name="微软雅黑")
    set_paragraph_runs(table.rows[2].cells[0].paragraphs[2], ["{{ trend_assessment }}"], target_font_name="微软雅黑")

    set_paragraph_runs(table.rows[3].cells[0].paragraphs[0], ["{{ monitor_heading }}", "", "", "", "", ""])

    monitoring_cell = table.rows[4].cells[0]
    set_paragraph_runs(monitoring_cell.paragraphs[0], ["1.", "\u00a0 ", "{{ monitor_subheading }}"])
    set_paragraph_runs(monitoring_cell.paragraphs[1], ["{{ waf_detail }}", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[4], ["\u00a0", "{{ waf_qps_detail }}", "", "", "", "", "", "", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[7], ["{{ cfw_detail }}", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[10], ["{{ cfw_bandwidth_detail }}", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[13], ["{{ hss_detail }}", "", "", "", "", "", "", "", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[17], ["\u00a0", "{{ ddos_detail }}", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[20], ["\u00a0", "{{ secmaster_detail }}", "", "", "", "", "", "", "", "", "", "", "", "", ""], target_font_name="微软雅黑")
    set_paragraph_runs(monitoring_cell.paragraphs[23], ["2.", "\u00a0 ", "{{ emergency_heading }}"])
    set_paragraph_runs(monitoring_cell.paragraphs[24], ["{{ emergency_response }}"], target_font_name="微软雅黑")

    set_paragraph_runs(table.rows[6].cells[0].paragraphs[0], ["{{ attack_path_assessment }}"], target_font_name="微软雅黑")
    set_paragraph_runs(table.rows[8].cells[0].paragraphs[0], ["{{ key_work_content }}"], target_font_name="微软雅黑")

    operator_table = table.rows[10].cells[0].tables[0]
    set_paragraph_runs(operator_table.rows[1].cells[0].paragraphs[0], ["{{ operator_group }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[1].cells[1].paragraphs[0], ["{{ operator_1_name }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[1].cells[2].paragraphs[0], ["{{ operator_1_phone }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[1].cells[3].paragraphs[0], ["{{ operator_1_role }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[1].cells[4].paragraphs[0], ["{{ operator_shared_responsibility }}"], target_font_name="微软雅黑")

    set_paragraph_runs(operator_table.rows[2].cells[1].paragraphs[0], ["{{ operator_2_name }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[2].cells[2].paragraphs[0], ["{{ operator_2_phone }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[2].cells[3].paragraphs[0], ["{{ operator_2_role }}"], target_font_name="微软雅黑")

    set_paragraph_runs(operator_table.rows[3].cells[1].paragraphs[0], ["{{ operator_3_name }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[3].cells[2].paragraphs[0], ["{{ operator_3_phone }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[3].cells[3].paragraphs[0], ["{{ operator_3_role }}", ""], target_font_name="微软雅黑")

    set_paragraph_runs(operator_table.rows[4].cells[1].paragraphs[0], ["{{ operator_4_name }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[4].cells[2].paragraphs[0], ["{{ operator_4_phone }}"], target_font_name="微软雅黑")
    set_paragraph_runs(operator_table.rows[4].cells[3].paragraphs[0], ["{{ operator_4_role }}", ""], target_font_name="微软雅黑")

    doc.save(TARGET_DOCX)
    print(TARGET_DOCX)


if __name__ == "__main__":
    main()
