#!/usr/bin/env python3
"""解析日报 DOCX 并保存为 daily_security_reports 的导入脚本

用法示例：
  python scripts/import_docx_report.py --file path/to/日报.docx --report-date 2026-05-20 --operator 导入者

脚本会尝试从 DOCX 中提取：`business_stability`、`trend_comparison`、`overall_assessment`，以及摘要句中的若干数值字段（WAF/CFW/HSS/DDOS/SecMaster），并调用 `app.db.save_daily_report` 保存。
"""
from __future__ import annotations

import argparse
import re
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict

from docx import Document

from app.db import save_daily_report, get_daily_report_by_date


def text_of_paragraph(p):
    return (p.text or "").strip()


def extract_summary_numbers(text: str) -> Dict[str, int]:
    # 默认都为 0
    out = {
        "waf_attacks": 0,
        "waf_blocked": 0,
        "waf_ips_banned": 0,
        "cfw_attacks": 0,
        "cfw_unblocked": 0,
        "hss_alerts": 0,
        "hss_unclosed_event_count": 0,
        "ddos_cleanings": 0,
        "ddos_blackholes": 0,
        "secmaster_alerts": 0,
        "secmaster_unclosed_event_count": 0,
    }

    def first_int(pattern):
        m = re.search(pattern, text)
        return int(m.group(1)) if m else 0

    out["waf_attacks"] = first_int(r"WAF(?:应用防火墙)?遭受攻击(\d+)次")
    out["waf_blocked"] = first_int(r"拦截攻击告警(\d+)次")
    out["waf_ips_banned"] = first_int(r"已对(\d+)个IP进行汇报封禁处理")
    out["cfw_attacks"] = first_int(r"CFW(?:云防火墙)?遭受攻击(\d+)次")
    out["cfw_unblocked"] = first_int(r"(\d+)次攻击未被拦截")
    out["hss_alerts"] = first_int(r"HSS(?:主机安全)?发生告警(\d+)次")
    # 两处未闭环事件，先匹配所有出现的数字并分配
    unclosed = re.findall(r"未闭环事件(\d+)次", text)
    if unclosed:
        if len(unclosed) >= 1:
            out["hss_unclosed_event_count"] = int(unclosed[0])
        if len(unclosed) >= 2:
            out["secmaster_unclosed_event_count"] = int(unclosed[1])

    out["ddos_cleanings"] = first_int(r"清洗(\d+)次")
    ddos_bh = re.search(r"(\d+)黑洞", text)
    if ddos_bh:
        out["ddos_blackholes"] = int(ddos_bh.group(1))

    out["secmaster_alerts"] = first_int(r"SecMaster(?:态势感知)?检测告警(\d+)次")

    return out


def split_trend_and_overall(text: str) -> (str, str):
    # 如果同时包含总体来看，则以该词为断点
    idx = text.find("总体来看，")
    if idx != -1:
        trend = text[:idx].strip()
        overall = text[idx:].strip()
        return trend, overall
    # 无总体判断，尝试以句号分割前后
    parts = re.split(r"(?<=。)\s*", text)
    if len(parts) >= 2:
        return parts[0].strip(), "".join(parts[1:]).strip()
    return text.strip(), ""


def _collect_texts(doc: Document) -> list[str]:
    """Collect all text blocks from both paragraphs and table cells."""
    texts: list[str] = []
    for p in doc.paragraphs:
        t = text_of_paragraph(p)
        if t:
            texts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    # Cell text may contain multiple paragraphs separated by \n
                    for line in t.split("\n"):
                        line = line.strip()
                        if line:
                            texts.append(line)
    return texts


def _find_section(texts: list[str], marker: str) -> int | None:
    for i, t in enumerate(texts):
        if t.startswith(marker):
            return i
    return None


def _first_int(pattern: str, text: str) -> int:
    m = re.search(pattern, text)
    return int(m.group(1)) if m else 0


def _normalize_nbsp(text: str) -> str:
    return text.replace("\xa0", " ").replace(" ", " ")


def parse_docx(path: Path) -> Dict[str, object]:
    doc = Document(str(path))
    texts = _collect_texts(doc)
    # Normalize non-breaking spaces
    texts = [_normalize_nbsp(t) for t in texts]

    payload: Dict[str, object] = {}

    # --- Section 1: 总体安全态势 ---
    idx1 = _find_section(texts, "一、总体安全态势")
    if idx1 is not None:
        # Collect all text between section 1 heading and section 2 heading
        idx2_start = _find_section(texts, "二、安全监控")
        if idx2_start is None:
            idx2_start = len(texts)
        s1_blocks = texts[idx1 + 1 : idx2_start]
        s1_text = "\n".join(s1_blocks)

        # First line is business_stability
        payload["business_stability"] = s1_blocks[0] if s1_blocks else ""

        # Extract summary numbers from section 1 text
        nums = extract_summary_numbers(s1_text)
        payload.update(nums)

        # Trend/overall are in the last block (skip business_stability and summary lines)
        trend_block = s1_blocks[-1] if len(s1_blocks) >= 2 else ""
        trend, overall = split_trend_and_overall(trend_block)
        payload["trend_comparison"] = trend or "较昨日趋势相比较，因缺少基线数据，暂无法开展趋势对比。"
        payload["overall_assessment"] = overall or "总体来看，整体安全态势保持平稳可控。"

    # --- Section 2: 安全监控 ---
    idx2 = None
    for i, t in enumerate(texts):
        if t.startswith("二、安全监控"):
            idx2 = i
            break

    if idx2 is not None:
        # Parse monitor time from heading
        m = re.search(r"监控时间[：:]\s*(.+?)\s*[~～]\s*(.+?)[）\)]?\s*$", texts[idx2])
        if m:
            payload["monitor_start"] = m.group(1).strip()
            payload["monitor_end"] = m.group(2).strip()

        # Collect all text from section 2 onwards (until section 3)
        s2_end = _find_section(texts, "三、攻击路径评估")
        if s2_end is None:
            s2_end = len(texts)
        s2_text = "\n".join(texts[idx2 + 1 : s2_end])

        # WAF detail
        payload["waf_detail_attacks"] = _first_int(r"WAF(?:应用防火墙)?遭受攻击(\d+)次", s2_text)
        payload["waf_detail_blocked"] = _first_int(r"拦截攻击告警(\d+)", s2_text)
        payload["waf_ips_banned"] = _first_int(r"已对(\d+)个IP进行汇报封禁处理", s2_text)
        waf_qps_match = re.search(r"QPS的规格为(.+?)，监测", s2_text)
        if waf_qps_match:
            payload["waf_qps_specs"] = waf_qps_match.group(1).strip()[:128]
        waf_peak_range = re.search(r"最大值时间段为([^，,]+)", s2_text)
        if waf_peak_range:
            payload["waf_qps_peak_range"] = waf_peak_range.group(1).strip()[:64]
        waf_peak = re.search(r"峰值为(\d+)", s2_text)
        if waf_peak:
            payload["waf_qps_peak_value"] = int(waf_peak.group(1))
        payload["waf_exceeded_spec"] = 1 if "已超出规格" in s2_text else 0

        # CFW detail
        payload["cfw_detail_attacks"] = _first_int(r"CFW(?:云防火墙)?遭受攻击(\d+)次攻击", s2_text)
        payload["cfw_detail_unblocked"] = _first_int(r"(\d+)次攻击未被拦截", s2_text)
        cfw_bw = re.search(r"防护带宽为([^，。]+)", s2_text)
        if cfw_bw:
            payload["cfw_bandwidth_spec"] = cfw_bw.group(1).strip()[:128]
        cfw_peak_range = re.search(r"带宽峰值时间段为([^，。]+(?:，[^，。]+)?)", s2_text)
        if cfw_peak_range:
            payload["cfw_peak_inbound_range"] = cfw_peak_range.group(1).strip()[:64]
        cfw_in_peak = re.search(r"入方向流量峰值([^，,]+)", s2_text)
        if cfw_in_peak:
            payload["cfw_inbound_peak"] = cfw_in_peak.group(1).strip()[:64]
        cfw_in_95 = re.search(r"入方向95带宽值([^，,。]+)", s2_text)
        if cfw_in_95:
            payload["cfw_inbound_95th"] = cfw_in_95.group(1).strip()[:64]
        payload["cfw_exceeded_spec"] = 1 if re.search(r"入方向流量峰值.*入方向95带宽.*超出规格", s2_text) else 0

        # HSS detail
        hss_match = re.search(
            r"HSS(?:主机安全)?发生告警(\d+)次[，,]\s*其中致命告警(\d+)次[，,]\s*高危告警(\d+)次[，,]\s*中危告警(\d+)次[，,]\s*低危告警(\d+)次",
            s2_text,
        )
        if hss_match:
            payload["hss_detail_total"] = int(hss_match.group(1))
            payload["hss_detail_fatal"] = int(hss_match.group(2))
            payload["hss_detail_high"] = int(hss_match.group(3))
            payload["hss_detail_medium"] = int(hss_match.group(4))
            payload["hss_detail_low"] = int(hss_match.group(5))
        hss_closed = re.search(r"((?:无成功入侵事件，)?已全部闭环)", s2_text)
        if hss_closed:
            payload["hss_closed_loop_status"] = hss_closed.group(1)

        # DDoS detail
        ddos_text = s2_text
        payload["ddos_detail_cleanings"] = _first_int(r"IP清洗(\d+)次", ddos_text)
        ddos_bh = re.search(r"(\d+)\s*黑洞", ddos_text)
        if ddos_bh:
            payload["ddos_detail_blackholes"] = int(ddos_bh.group(1)) if "无黑洞" not in ddos_text else 0

        # SecMaster detail
        sm_match = re.search(
            r"SecMaster(?:态势感知)?发生告警(\d+)次[，,]\s*其中致命告警(\d+)次[，,]\s*高危告警(\d+)次[，,]\s*中危告警(\d+)次[，,]\s*低危告警(\d+)次[，,]\s*提示告警(\d+)次",
            s2_text,
        )
        if sm_match:
            payload["secmaster_detail_total"] = int(sm_match.group(1))
            payload["secmaster_detail_fatal"] = int(sm_match.group(2))
            payload["secmaster_detail_high"] = int(sm_match.group(3))
            payload["secmaster_detail_medium"] = int(sm_match.group(4))
            payload["secmaster_detail_low"] = int(sm_match.group(5))
            payload["secmaster_detail_info"] = int(sm_match.group(6))

        # Emergency response
        er_match = re.search(r"2[\.\s]+事件应急响应[：:]\s*\n?(.*?)(?=\n(?:三、|图\d|$))", s2_text, re.DOTALL)
        if er_match:
            payload["emergency_response"] = er_match.group(1).strip()
        else:
            # Try simpler pattern
            er_simple = re.search(r"事件应急响应[：:]\s*(.+)", s2_text)
            if er_simple:
                payload["emergency_response"] = er_simple.group(1).strip()

    # --- Section 3: 攻击路径评估 ---
    idx3 = _find_section(texts, "三、攻击路径评估")
    if idx3 is not None and idx3 + 1 < len(texts):
        next_text = texts[idx3 + 1]
        if not next_text.startswith("四、"):
            payload["attack_path_assessment"] = next_text

    # --- Section 4: 重点工作内容 ---
    idx4 = _find_section(texts, "四、重点工作内容")
    if idx4 is not None and idx4 + 1 < len(texts):
        next_text = texts[idx4 + 1]
        if not next_text.startswith("五、"):
            payload["key_work_content"] = next_text

    # Fill defaults for fields not present
    payload.setdefault("attack_path_assessment", "暂无")
    payload.setdefault("key_work_content", "暂无")
    payload.setdefault("emergency_response", "无。")

    return payload


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", "-f", required=True, help="DOCX 文件路径")
    parser.add_argument("--report-date", "-d", default=None, help="报告日期，格式 YYYY-MM-DD，默认昨天")
    parser.add_argument("--operator", "-o", default="导入者", help="操作人显示名称")
    parser.add_argument("--preview", action="store_true", help="只打印解析结果，不写入数据库")

    args = parser.parse_args()
    path = Path(args.file)
    if not path.exists():
        print(f"文件不存在: {path}")
        return

    if args.report_date:
        report_date = args.report_date
    else:
        yesterday = datetime.now() - timedelta(days=1)
        report_date = yesterday.strftime("%Y-%m-%d")

    payload = parse_docx(path)

    print("解析结果：")
    for k, v in payload.items():
        print(f"  {k}: {v}")

    if args.preview:
        return

    # 读取现有并合并，以保持缺失字段为空串或 0
    existing = get_daily_report_by_date(report_date) or {}
    # normalize keys expected by save_daily_report: many fields default to 0/''
    merged = dict(existing)
    merged.update(payload)

    save_daily_report(report_date, merged, args.operator)
    print(f"已保存到数据库，report_date={report_date}")


if __name__ == "__main__":
    main()
